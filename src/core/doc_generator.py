import google.generativeai as genai
import os
from dotenv import load_dotenv
from playwright.sync_api import sync_playwright
from PIL import Image
import io
from docx import Document
from docx.shared import Inches

load_dotenv()

# --- Initialize Gemini API ---
try:
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        raise ValueError("GEMINI_API_KEY not found in .env file.")
    genai.configure(api_key=gemini_api_key)
    print("Gemini API configured successfully.")
except Exception as e:
    print(f"CRITICAL ERROR: Failed to configure Gemini API. Ensure GEMINI_API_KEY is correct in .env. Details: {e}")
    gemini_api_key = None 

# Use a suitable Gemini model for contextual analysis and text generation
text_model = genai.GenerativeModel('gemini-1.5-flash')

# --- Web Automation and Screenshot Function ---
def take_screenshot(url: str, selector: str = None) -> bytes:
    """Navigates to a URL and takes a screenshot of a specific selector or the whole page."""
    image_bytes = None
    with sync_playwright() as p:
        try:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(url)
            page.wait_for_load_state("networkidle") # Wait for all network activity to cease

            if selector:
                # Take screenshot of a specific element
                element = page.locator(selector)
                image_bytes = element.screenshot()
            else:
                # Take full page screenshot
                image_bytes = page.screenshot(full_page=True)
            
            browser.close()
        except Exception as e:
            print(f"ERROR: Playwright failed to take screenshot. Details: {e}")
            raise Exception(f"Failed to automate browser: {e}. Check URL or selector.")
    return image_bytes

# --- AI Description Generator ---
def generate_description(image_bytes: bytes, context: str) -> str:
    """Generates a text description for a screenshot using Gemini Vision AI."""
    if not gemini_api_key:
        raise ValueError("Gemini API Key is not configured.")
        
    try:
        image = Image.open(io.BytesIO(image_bytes))
        prompt = f"""
        You are an expert technical writer and documentation specialist.
        Based on the provided screenshot and context, generate a clear and concise description.
        
        **Context:** {context}
        
        **Task:**
        1.  Analyze the screenshot.
        2.  Explain the functionality or purpose of the web page/element shown.
        3.  Describe the key visual components and how they relate to the context.
        
        Provide the description in a professional tone, suitable for technical documentation.
        """
        response = text_model.generate_content([prompt, image])
        return response.text
    except Exception as e:
        raise Exception(f"Failed to generate description with Gemini: {e}")


# --- MS Word Document Generator ---
def create_documentation_report(data: list, output_filename: str):
    """
    Creates an MS Word document with screenshots and their descriptions.
    
    Args:
        data (list): A list of dictionaries, where each dict has 'description' and 'image_bytes'.
        output_filename (str): The name of the Word document to save.
    """
    doc = Document()
    doc.add_heading('Website Documentation Report', level=1)
    
    for item in data:
        doc.add_heading(item['context'], level=2)
        doc.add_paragraph(item['description'])
        
        image_stream = io.BytesIO(item['image_bytes'])
        doc.add_picture(image_stream, width=Inches(6)) # Add image with a fixed width
        doc.add_paragraph() # Add space after image

    doc.save(output_filename)

    # To get the file bytes for Streamlit download
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream.getvalue()

# Example Usage (typically run via app.py)
if __name__ == "__main__":
    print("This module contains the core logic for the Documentation Generator.")
    print("Please run app.py to interact with this application.")